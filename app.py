import streamlit as st
import sqlite3
from datetime import datetime
import os

# =====================================
# مكتبات تصدير التقارير
# =================================
from io import BytesIO

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph
)

from reportlab.lib.styles import getSampleStyleSheet
# ===========================
# =====================================
# إنشاء تقرير Word
# =====================================

def create_word_report(
    report_title,
    office,
    lawyer,
    from_date,
    to_date,
    headers,
    rows
):

    file = BytesIO()
    doc = Document()

    # =================================
    # إعداد الصفحة
    # =================================

    section = doc.sections[0]

    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width, section.page_height = (
        section.page_height,
        section.page_width
    )

    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.left_margin = Cm(1.2)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    style.font.size = Pt(11)

    COLOR_BLUE = RGBColor(17, 52, 120)
    COLOR_GOLD = RGBColor(184, 134, 11)
    COLOR_BLACK = RGBColor(0, 0, 0)

    # =================================
    # رأس التقرير
    # =================================

    head = doc.add_table(rows=1, cols=3)
    head.alignment = WD_TABLE_ALIGNMENT.CENTER

    head.columns[0].width = Cm(2.5)
    head.columns[1].width = Cm(18)
    head.columns[2].width = Cm(3)

    for cell in head.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")

        for edge in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "nil")
            tcBorders.append(border)

        tcPr.append(tcBorders)

    # الميزان

    p = head.cell(0, 0).paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    run = p.add_run("⚖")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = COLOR_GOLD

    # عنوان الهيئة

    p = head.cell(0, 1).paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    r = p.add_run("الهيئة القومية للتأمين الاجتماعى\n")
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = COLOR_BLUE

    r = p.add_run("الإدارة المركزية للشئون القانونية\n")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = COLOR_BLUE

    r = p.add_run("الإدارة العامة للقضايــــــــــــا\n")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = COLOR_BLUE

    r = p.add_run(f"ديوان عام منطقة : {office}")
    r.bold = True
    r.font.size = Pt(12)

    # شعار الهيئة

    p = head.cell(0, 2).paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    try:
        run = p.add_run()
        run.add_picture("logo.png", width=Mm(18))
    except:
        pass

    doc.add_paragraph()

    # =================================
    # عنوان التقرير
    # =================================

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    r = p.add_run(report_title)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = COLOR_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    r = p.add_run(
        f"خلال الفترة من {from_date.strftime('%d/%m/%Y')} حتى {to_date.strftime('%d/%m/%Y')}"
    )
    r.bold = True
    r.font.size = Pt(12)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    r = p.add_run(f"طرف الأستاذ / {lawyer} (المحامى)")
    r.bold = True
    r.font.size = Pt(12)

    doc.add_paragraph()

    # =================================
        # =================================
    # ختام التقرير
    # =================================

    doc.add_paragraph()

    p = doc.add_paragraph()

    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    r = p.add_run("وتفضلوا بقبول وافر الاحترام")

    r.bold = True

    r.font.size = Pt(13)

    doc.add_paragraph()
    doc.add_paragraph()

    # =================================
# =================================
# التوقيعات
# =================================

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

r = p.add_run("وتفضلوا بقبول وافر الاحترام")
r.bold = True
r.font.size = Pt(12)

doc.add_paragraph()

sign = doc.add_table(rows=2, cols=3)
sign.alignment = WD_TABLE_ALIGNMENT.CENTER

# إزالة الحدود
for row in sign.rows:
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "nil")
            tcBorders.append(e)
        tcPr.append(tcBorders)

# الصف الأول
sign.cell(0,0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
sign.cell(0,0).text = "عضو الإدارة"

sign.cell(0,1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
sign.cell(0,1).text = "مدير عام الإدارات القانونية"

sign.cell(0,2).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
sign.cell(0,2).text = "مدير الإدارة"

# الصف الثانى
sign.cell(1,0).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
sign.cell(1,0).text = "......................."

sign.cell(1,1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
sign.cell(1,1).text = "......................."

sign.cell(1,2).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
sign.cell(1,2).text = "......................."

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

r = p.add_run(
    f"تحريراً فى : {datetime.now().strftime('%d/%m/%Y')}"
)
r.bold = True
    # =================================
    # التاريخ
    # =================================

    p = doc.add_paragraph()

    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    p.add_run(
        f"تحريراً فى : {datetime.now().strftime('%d/%m/%Y')}"
    ).bold = True

    # =================================
    # ترقيم الصفحات
    # =================================

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    footer = section.footer

    fp = footer.paragraphs[0]

    fp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    fp.add_run("صفحة ")

    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)

    fp.add_run(" من ")

    fld2 = OxmlElement("w:fldSimple")
    fld2.set(qn("w:instr"), "NUMPAGES")
    fp._p.append(fld2)

    # =================================
    # حفظ الملف
    # =================================

    doc.save(file)

    file.seek(0)

    return file

# =====================================
# نهاية دالة التقرير
# =====================================
# =====================================
# =====================================
# إعداد الصفحة
# =====================================

st.set_page_config(
    page_title="إدارة القضايا",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# =====================================
# Session State
# =====================================

if "page" not in st.session_state:
    st.session_state.page = "home"

if "selected_case" not in st.session_state:
    st.session_state.selected_case = None

# =====================================
# قاعدة البيانات
# =====================================

conn = sqlite3.connect(
    "cases.db",
    check_same_thread=False
)

cur = conn.cursor()

# =====================================
# جدول القضايا
# =====================================

cur.execute("""
CREATE TABLE IF NOT EXISTS cases(

id INTEGER PRIMARY KEY AUTOINCREMENT,

case_type TEXT,
court_type TEXT,
court_name TEXT,
mission TEXT,

case_number TEXT,
judicial_year TEXT,
circuit TEXT,
case_category TEXT,

plaintiff TEXT,
defendant TEXT,

subject TEXT,

notes TEXT,

whatsapp_enabled INTEGER,
whatsapp_number TEXT,

created_at TEXT

)
""")

# =====================================
# جدول الجلسات
# =====================================

cur.execute("""

CREATE TABLE IF NOT EXISTS sessions(

id INTEGER PRIMARY KEY AUTOINCREMENT,

case_id INTEGER,

session_date TEXT,

roll_number TEXT,

procedure TEXT,

created_at TEXT

)

""")

# =====================================
# جدول المستندات
# =====================================

cur.execute("""

CREATE TABLE IF NOT EXISTS documents(

id INTEGER PRIMARY KEY AUTOINCREMENT,

case_id INTEGER,

document_type TEXT,

document_description TEXT,

file_name TEXT,

file_path TEXT,

uploaded_at TEXT

)

""")

conn.commit()
# =====================================
# التصميم العام (CSS)
# =====================================

st.markdown("""
<style>

#MainMenu{visibility:hidden;}
header{visibility:hidden;}
footer{visibility:hidden;}

html,body,[class*="css"]{
    direction:rtl;
}

/* الخلفية */

.stApp{

background:
linear-gradient(
180deg,
#5A3A22 0%,
#4A2F1C 35%,
#352114 70%,
#24160D 100%
);

}

/* ميزان الخلفية */

.watermark{

position:fixed;

top:50%;

left:50%;

transform:translate(-50%,-50%);

font-size:320px;

opacity:.05;

z-index:0;

}

/* رأس البرنامج */

.header{

text-align:center;

position:relative;

z-index:999;

}

.logo{

font-size:90px;

margin-bottom:-15px;

}

.title1{

font-size:34px;

font-weight:bold;

color:#FFD700;

}

.title2{

font-size:27px;

font-weight:bold;

color:white;

}

.title3{

font-size:27px;

font-weight:bold;

color:white;

}

.title4{

font-size:27px;

font-weight:bold;

color:white;

}

/* الأزرار */

.stButton>button{

width:100%;

height:85px;

font-size:24px;

font-weight:bold;

border-radius:18px;

background:#6D4C41;

color:white;

border:2px solid gold;

transition:.3s;

}

.stButton>button:hover{

background:#8D6E63;

transform:scale(1.02);

box-shadow:0 0 18px gold;

}

/* الكتابة */

label,p,h1,h2,h3,h4,h5,h6,span{

color:white !important;

font-weight:bold;

}

/* الشريط السفلى */

.news-bar{

position:fixed;

bottom:0;

right:0;

width:100%;

height:45px;

background:#1A120B;

border-top:2px solid gold;

overflow:hidden;

z-index:999999;

}

.news-text{

position:absolute;

right:-100%;

white-space:nowrap;

line-height:45px;

font-size:20px;

font-weight:bold;

animation:scrollText 35s linear infinite;

}

@keyframes scrollText{

from{

right:-100%;

}

to{

right:100%;

}

}

</style>

""",unsafe_allow_html=True)

# =====================================
# الميزان بالخلفية
# =====================================

st.markdown("""

<div class="watermark">

⚖️

</div>

""",unsafe_allow_html=True)

# =====================================
# رأس البرنامج
# =====================================

st.markdown("""

<div class="header">

<div class="logo">
⚖️
</div>

<div class="title1">
           الهيئة القومية للتأمين الاجتماعى
</div>

<div class="title2">
الإدارة المركزية للشئون القانونية
</div>

<div class="title3">
الإدارة العامة للقضايا
</div>

<div class="title4">
ديوان عام منطقة البحيرة
</div>

</div>

<br>

""",unsafe_allow_html=True)

# =====================================
# الشريط السفلى
# =====================================

st.markdown("""

<div class="news-bar">

<div class="news-text">

<span style="color:#FFD700;">
✨ مع تحيات / وليد حماد
</span>

<span style="color:white;">
&nbsp;&nbsp;|&nbsp;&nbsp;
</span>

<span style="color:#00FFFF;">
الإدارة العامة للقضايا
</span>

<span style="color:white;">
&nbsp;&nbsp;|&nbsp;&nbsp;
</span>

<span style="color:#7CFC00;">
الإدارة المركزية للشئون القانونية
</span>

<span style="color:white;">
&nbsp;&nbsp;|&nbsp;&nbsp;
</span>

<span style="color:#FFA500;">
ديوان عام منطقة البحيرة
</span>

<span style="color:white;">
&nbsp;&nbsp;|&nbsp;&nbsp;
</span>

<span style="color:#FF4444;">
الهيئة القومية للتأمين الاجتماعى
</span>

</div>

</div>

""",unsafe_allow_html=True)
# =====================================
# الصفحة الرئيسية
# =====================================

if st.session_state.page == "home":

    st.markdown("<br>", unsafe_allow_html=True)

    c1, c2, c3 = st.columns([1,3,1])

    with c2:

        if st.button(
            "⚖️ تسجيل القضايا",
            use_container_width=True
        ):
            st.session_state.page = "cases"
            st.rerun()

        st.markdown("<br>", unsafe_allow_html=True)

        if st.button(
            "📋 الحصر العام",
            use_container_width=True
        ):
            st.session_state.page = "inventory"
            st.rerun()

        st.markdown("<br>", unsafe_allow_html=True)

        if st.button(
            "🔔 التنبيهات",
            use_container_width=True
        ):
            st.session_state.page = "notifications"
            st.rerun()

        st.markdown("<br>", unsafe_allow_html=True)

        if st.button(
            "📊 التقارير",
            use_container_width=True
        ):
            st.session_state.page = "reports"
            st.rerun()

        st.markdown("<br>", unsafe_allow_html=True)

        if st.button(
            "🗄️ الأرشيف",
            use_container_width=True
        ):
            st.session_state.page = "archive"
            st.rerun()

        st.markdown("<br>", unsafe_allow_html=True)

        if st.button(
            "📚 المكتبة القانونية",
            use_container_width=True
        ):
            st.session_state.page = "library"
            st.rerun()

        st.markdown("<br>", unsafe_allow_html=True)

        if st.button(
            "🔍 البحث عن دعوى",
            use_container_width=True
        ):
            st.session_state.page = "search"
            st.rerun()


# =====================================
# صفحات تحت الإنشاء
# =====================================

elif st.session_state.page == "notifications":

    st.info("🚧 سيتم تنفيذ قسم التنبيهات فى المرحلة القادمة.")

    if st.button("⬅ العودة للرئيسية"):
        st.session_state.page = "home"
        st.rerun()

elif st.session_state.page == "archive":

    st.info("🚧 سيتم تنفيذ قسم الأرشيف فى المرحلة القادمة.")

    if st.button("⬅ العودة للرئيسية"):
        st.session_state.page = "home"
        st.rerun()


elif st.session_state.page == "library":

    st.info("🚧 سيتم تنفيذ قسم المكتبة القانونية فى المرحلة القادمة.")

    if st.button("⬅ العودة للرئيسية"):
        st.session_state.page = "home"
        st.rerun()


elif st.session_state.page == "search":

    st.info("🚧 سيتم تنفيذ البحث عن الدعاوى فى المرحلة القادمة.")

    if st.button("⬅ العودة للرئيسية"):
        st.session_state.page = "home"
        st.rerun()
        # ==========================================================
# تسجيل القضايا - الجزء الأول
# ==========================================================

elif st.session_state.page == "cases":

    st.markdown("""

    <div style="
    background:#4E342E;
    padding:15px;
    border-radius:15px;
    border:2px solid gold;
    text-align:center;
    color:white;
    font-size:30px;
    font-weight:bold;
    box-shadow:0 0 15px rgba(255,215,0,.4);
    ">

    ⚖️ تسجيل قضية جديدة

    </div>

    """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    col1, col2 = st.columns(2)

    # =====================================
    # العمود الأول
    # =====================================

    with col1:

        case_type = st.selectbox(
            "نوع الدعوى",
            [
                "دعوى",
                "استئناف",
                "طعن"
            ]
        )

        court_type = st.selectbox(
            "المحكمة",
            [
                "الابتدائية",
                "الاستئناف",
                "النقض",
                "الإدارية",
                "القضاء الإدارى",
                "الإدارية العليا"
            ]
        )

        court_name = st.text_input(
            "اسم المحكمة"
        )

        mission = ""

        if case_type == "استئناف":

            mission = st.text_input(
                "المأمورية"
            )

        case_number = st.text_input(
            "رقم الدعوى / الاستئناف / الطعن"
        )

        judicial_year = st.text_input(
            "السنة القضائية"
        )

        circuit = st.text_input(
            "الدائرة"
        )

        case_category = st.text_input(
            "النوع"
        )
        # =====================================
    # العمود الثانى
    # =====================================

    with col2:

        plaintiff = st.text_input(
            "اسم المدعى / المستأنف / الطاعن"
        )

        defendant = st.text_input(
            "اسم المدعى عليه / المستأنف ضده / المطعون ضده"
        )

        subject = st.text_area(
            "موضوع الدعوى",
            height=130
        )

        first_session_date = st.date_input(
            "تاريخ أول جلسة"
        )

        roll_number = st.text_input(
            "الرول"
        )

        first_procedure = st.text_area(
            "سبب الجلسة",
            height=120
        )

        notes = st.text_area(
            "ملاحظات",
            height=120
        )

    st.markdown("<hr>", unsafe_allow_html=True)

    # =====================================
    # تنبيهات الواتساب
    # =====================================

    whatsapp_enabled = st.checkbox(
        "📲 تفعيل التنبيهات عبر واتساب"
    )

    whatsapp_number = ""

    if whatsapp_enabled:

        whatsapp_number = st.text_input(
            "رقم هاتف واتساب"
        )

    st.markdown("<br>", unsafe_allow_html=True)

    # =====================================
    # رفع المستندات
    # =====================================

    document_type = st.selectbox(

        "نوع المستند",

        [

            "صحيفة الدعوى",

            "صحيفة الاستئناف",

            "صحيفة الطعن"

        ]

    )

    uploaded_file = st.file_uploader(

        "تحميل المستند",

        type=[
            "pdf",
            "doc",
            "docx"
        ]

    )

    st.markdown("<br>", unsafe_allow_html=True)
    # =====================================
    # أزرار الحفظ والحذف
    # =====================================

    col_save, col_delete = st.columns(2)

    # =====================================
    # حفظ القضية
    # =====================================

    with col_save:

        if st.button(
            "💾 حفظ القضية",
            use_container_width=True
        ):

            # التحقق من البيانات الأساسية

            if case_number.strip() == "":
                st.error("برجاء إدخال رقم القضية")
                st.stop()

            if court_name.strip() == "":
                st.error("برجاء إدخال اسم المحكمة")
                st.stop()

            if plaintiff.strip() == "":
                st.error("برجاء إدخال اسم المدعى")
                st.stop()

            # حفظ القضية

            cur.execute("""

            INSERT INTO cases(

            case_type,
            court_type,
            court_name,
            mission,
            case_number,
            judicial_year,
            circuit,
            case_category,
            plaintiff,
            defendant,
            subject,
            notes,
            whatsapp_enabled,
            whatsapp_number,
            created_at

            )

            VALUES(

            ?,?,?,?,?,?,?,?,?,?,?,?,?,?,?

            )

            """,

            (

            case_type,
            court_type,
            court_name,
            mission,
            case_number,
            judicial_year,
            circuit,
            case_category,
            plaintiff,
            defendant,
            subject,
            notes,
            int(whatsapp_enabled),
            whatsapp_number,
            datetime.now().strftime(
                "%Y-%m-%d %H:%M:%S"
            )

            )

            )

            conn.commit()

            case_id = cur.lastrowid

            # =====================================
            # حفظ أول جلسة
            # =====================================

            cur.execute("""

            INSERT INTO sessions(

            case_id,
            session_date,
            roll_number,
            procedure,
            created_at

            )

            VALUES(

            ?,?,?,?,?

            )

            """,

            (

            case_id,
            str(first_session_date),
            roll_number,
            first_procedure,
            datetime.now().strftime(
                "%Y-%m-%d %H:%M:%S"
            )

            )

            )

            conn.commit()

            st.success("✅ تم حفظ القضية بنجاح")

            st.session_state.page = "inventory"

            st.rerun()

    # =====================================
    # حذف البيانات المدخلة
    # =====================================

    with col_delete:

        if st.button(
            "🗑 حذف البيانات",
            use_container_width=True
        ):

            st.rerun()
            # =====================================
            # إنشاء مجلد المستندات
            # =====================================

            if uploaded_file is not None:

                documents_folder = "Documents"

                if not os.path.exists(documents_folder):
                    os.mkdir(documents_folder)

                case_folder = os.path.join(
                    documents_folder,
                    f"{case_number}-{judicial_year}"
                )

                if not os.path.exists(case_folder):
                    os.mkdir(case_folder)

                file_path = os.path.join(
                    case_folder,
                    uploaded_file.name
                )

                with open(file_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())

                cur.execute("""

                INSERT INTO documents(

                case_id,
                document_type,
                document_description,
                file_name,
                file_path,
                uploaded_at

                )

                VALUES(

                ?,?,?,?,?,?

                )

                """,

                (

                case_id,

                document_type,

                f"{document_type} رقم {case_number}",

                uploaded_file.name,

                file_path,

                datetime.now().strftime(
                    "%Y-%m-%d %H:%M:%S"
                )

                )

                )

                conn.commit()

            st.success("✅ تم حفظ القضية وجميع البيانات بنجاح")

            st.balloons()

            st.session_state.page = "inventory"

            st.rerun()
            # =====================================
    # التحقق من رقم الواتساب
    # =====================================

    if whatsapp_enabled:

        if whatsapp_number.strip() != "":

            valid_prefix = (
                whatsapp_number.startswith("010") or
                whatsapp_number.startswith("011") or
                whatsapp_number.startswith("012") or
                whatsapp_number.startswith("015")
            )

            if len(whatsapp_number) != 11 or not valid_prefix:

                st.error(
                    "رقم واتساب غير صحيح (يجب أن يكون 11 رقم ويبدأ بـ 010 أو 011 أو 012 أو 015)"
                )
                st.stop()

    # =====================================
    # تحديد نوع المستند تلقائياً
    # =====================================

    if case_type == "دعوى":
        document_type = "صحيفة الدعوى"

    elif case_type == "استئناف":
        document_type = "صحيفة الاستئناف"

    else:
        document_type = "صحيفة الطعن"

    # =====================================
    # أزرار أسفل الصفحة
    # =====================================

    st.markdown("<br>", unsafe_allow_html=True)

    b1, b2 = st.columns(2)

    with b1:

        if st.button(
            "⬅ العودة للرئيسية",
            use_container_width=True,
            key="back_home_cases"
        ):

            st.session_state.page = "home"
            st.rerun()

    with b2:

        if st.button(
            "🧹 مسح جميع البيانات",
            use_container_width=True,
            key="clear_form"
        ):

            st.warning("تم مسح البيانات المدخلة.")

            st.rerun()
            # ==========================================================
# الحصر العام - الجزء الأول
# ==========================================================

elif st.session_state.page == "inventory":

    st.markdown("""
    <div style="
    background:#4E342E;
    border:2px solid gold;
    border-radius:15px;
    padding:15px;
    text-align:center;
    color:white;
    font-size:30px;
    font-weight:bold;
    box-shadow:0 0 15px rgba(255,215,0,.4);
    ">
    📋 الحصر العام للقضايا
    </div>
    """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # =====================================
    # الإحصائيات
    # =====================================

    cur.execute("SELECT COUNT(*) FROM cases")
    total_cases = cur.fetchone()[0]

    cur.execute("SELECT COUNT(*) FROM sessions")
    total_sessions = cur.fetchone()[0]

    c1, c2 = st.columns(2)

    with c1:

        st.metric(
            "⚖️ إجمالى القضايا",
            total_cases
        )

    with c2:

        st.metric(
            "📅 إجمالى الجلسات",
            total_sessions
        )

    st.markdown("---")

    # =====================================
    # البحث
    # =====================================

    search = st.text_input(
        "🔍 البحث فى القضايا"
    )

    # =====================================
    # الفلاتر
    # =====================================

    f1, f2 = st.columns(2)

    with f1:

        case_filter = st.selectbox(

            "نوع الدعوى",

            [

                "الكل",

                "دعوى",

                "استئناف",

                "طعن"

            ]

        )

    with f2:

        court_filter = st.text_input(
            "المحكمة"
    )
        # =====================================
    # قراءة القضايا
    # =====================================

    cur.execute("""

    SELECT

    c.id,
    c.case_type,
    c.case_number,
    c.judicial_year,
    c.circuit,
    c.case_category,
    c.court_name,
    c.mission,
    c.plaintiff,
    c.defendant,
    c.subject,

    (

        SELECT session_date

        FROM sessions s

        WHERE s.case_id=c.id

        ORDER BY session_date DESC

        LIMIT 1

    ) last_session,

    (

        SELECT procedure

        FROM sessions s

        WHERE s.case_id=c.id

        ORDER BY session_date DESC

        LIMIT 1

    ) last_procedure

    FROM cases c

    ORDER BY c.created_at DESC

    """)

    rows = cur.fetchall()

    # =====================================
    # تطبيق البحث والفلاتر
    # =====================================

    filtered = []

    for row in rows:

        text = " ".join(
            [str(x) if x else "" for x in row]
        ).lower()

        if search.strip():

            if search.lower() not in text:
                continue

        if case_filter != "الكل":

            if row[1] != case_filter:
                continue

        if court_filter.strip():

            if court_filter.lower() not in row[6].lower():
                continue

        filtered.append(row)

    st.markdown("<br>", unsafe_allow_html=True)

    if len(filtered) == 0:

        st.warning("لا توجد نتائج")

    else:

        st.success(f"عدد النتائج : {len(filtered)}")

        st.markdown("<br>", unsafe_allow_html=True)
        # =====================================
    # عرض القضايا
    # =====================================

    for row in filtered:

        case_id = row[0]

        plaintiff = (
            row[8].replace(
                "الهيئة القومية للتأمين الاجتماعى",
                "الهيئة"
            )
            if row[8] else ""
        )

        defendant = (
            row[9].replace(
                "الهيئة القومية للتأمين الاجتماعى",
                "الهيئة"
            )
            if row[9] else ""
        )

        st.markdown(f"""
        <div style="
        background:#4E342E;
        border:2px solid #D4AF37;
        border-radius:18px;
        padding:18px;
        margin-bottom:18px;
        box-shadow:0 0 10px rgba(212,175,55,.35);
        ">

        <div style="
        color:#FFD700;
        font-size:24px;
        font-weight:bold;
        ">

        ⚖️ {row[1]} رقم {row[2]} لسنة {row[3]}

        </div>

        <hr style="border:1px solid #8B7355;">

        <div style="color:white;font-size:18px;">

        🏛 <b>المحكمة :</b> {row[6]}

        </div>

        <div style="color:white;font-size:18px;">

        ⚖️ <b>الدائرة :</b> {row[4]}

        </div>

        <div style="color:white;font-size:18px;">

        📂 <b>النوع :</b> {row[5]}

        </div>

        <div style="color:white;font-size:18px;">

        👤 <b>الخصوم :</b><br>

        {plaintiff}<br>

        ضد<br>

        {defendant}

        </div>

        <div style="color:#00FFFF;font-size:18px;">

        📄 <b>الموضوع :</b><br>

        {row[10]}

        </div>

        <div style="color:#7CFC00;font-size:18px;">

        📅 <b>آخر جلسة :</b>

        {row[11] if row[11] else "لا يوجد"}

        </div>

        <div style="color:#FFA500;font-size:18px;">

        ⚖️ <b>آخر إجراء :</b><br>

        {row[12] if row[12] else "لا يوجد"}

        </div>

        </div>

        """, unsafe_allow_html=True)

        b1, b2, b3 = st.columns(3)

        with b1:

            if st.button(
                "📂 فتح القضية",
                key=f"open_{case_id}",
                use_container_width=True
            ):

                st.session_state.selected_case = case_id
                st.session_state.page = "case_details"
                st.rerun()

        with b2:

            if st.button(
                "✏️ تعديل",
                key=f"edit_{case_id}",
                use_container_width=True
            ):

                st.session_state.selected_case = case_id
                st.session_state.page = "edit_case"
                st.rerun()

        with b3:

            if st.button(
                "🗑 حذف",
                key=f"delete_{case_id}",
                use_container_width=True
            ):

                st.session_state.delete_case_id = case_id

        st.markdown("<br>", unsafe_allow_html=True)
        # =====================================
    # تأكيد حذف القضية
    # =====================================

    if "delete_case_id" not in st.session_state:
        st.session_state.delete_case_id = None

    if st.session_state.delete_case_id:

        st.markdown("---")

        st.error("⚠️ هل تريد حذف القضية نهائياً؟")

        d1, d2 = st.columns(2)

        with d1:

            if st.button(
                "✅ نعم - حذف القضية",
                use_container_width=True,
                key="confirm_delete_case"
            ):

                delete_id = st.session_state.delete_case_id

                cur.execute(
                    "DELETE FROM sessions WHERE case_id=?",
                    (delete_id,)
                )

                cur.execute(
                    "DELETE FROM documents WHERE case_id=?",
                    (delete_id,)
                )

                cur.execute(
                    "DELETE FROM cases WHERE id=?",
                    (delete_id,)
                )

                conn.commit()

                st.success("تم حذف القضية بنجاح")

                st.session_state.delete_case_id = None

                st.rerun()

        with d2:

            if st.button(
                "❌ إلغاء",
                use_container_width=True,
                key="cancel_delete_case"
            ):

                st.session_state.delete_case_id = None

                st.rerun()

    st.markdown("<br>", unsafe_allow_html=True)

    # =====================================
    # العودة للرئيسية
    # =====================================

    if st.button(
        "⬅ العودة للرئيسية",
        use_container_width=True,
        key="inventory_home"
    ):

        st.session_state.page = "home"

        st.rerun()
        # ==========================================================
# ملف القضية - الجزء الأول
# ==========================================================

elif st.session_state.page == "case_details":

    case_id = st.session_state.selected_case

    cur.execute("""

    SELECT *

    FROM cases

    WHERE id=?

    """,(case_id,))

    case = cur.fetchone()

    if not case:

        st.error("❌ القضية غير موجودة")

        if st.button("⬅ العودة للحصر العام"):

            st.session_state.page="inventory"

            st.rerun()

    else:

        st.markdown("""

        <div style="
        background:#4E342E;
        border:2px solid gold;
        border-radius:18px;
        padding:18px;
        text-align:center;
        color:white;
        font-size:32px;
        font-weight:bold;
        ">

        ⚖️ ملف القضية

        </div>

        """,unsafe_allow_html=True)

        st.markdown("<br>",unsafe_allow_html=True)

        # =====================================
        # بيانات القضية
        # =====================================

        st.markdown("""
        <h3 style='color:#FFD700'>
        📋 بيانات القضية
        </h3>
        """,unsafe_allow_html=True)

        t1,t2=st.columns(2)

        with t1:

            st.info(f"""
⚖️ نوع الدعوى

{case[1]}
""")

            st.info(f"""
🏛 المحكمة

{case[3]}
""")

            st.info(f"""
📄 رقم الدعوى

{case[5]}
""")

            st.info(f"""
📅 السنة القضائية

{case[6]}
""")

            st.info(f"""
⚖️ الدائرة

{case[7]}
""")

            if case[4]:

                st.info(f"""
📍 المأمورية

{case[4]}
""")

        with t2:

            st.success(f"""
👤 المدعى

{case[9]}
""")

            st.error(f"""
👤 المدعى عليه

{case[10]}
""")

            st.warning(f"""
📑 موضوع الدعوى

{case[11]}
""")

            if case[12]:

                st.info(f"""
📝 الملاحظات

{case[12]}
""")

        st.markdown("<hr>",unsafe_allow_html=True)

        st.markdown("""
        <h3 style='color:#FFD700'>
        📅 جدول الجلسات
        </h3>
        """,unsafe_allow_html=True)
        # =====================================
        # قراءة الجلسات
        # =====================================

        cur.execute("""

        SELECT

        id,
        roll_number,
        session_date,
        procedure

        FROM sessions

        WHERE case_id=?

        ORDER BY session_date ASC

        """,(case_id,))

        sessions=cur.fetchall()

        if sessions:

            st.markdown("""
            <style>

            table{
                width:100%;
                border-collapse:collapse;
            }

            th{
                background:#5D4037;
                color:#FFD700;
                text-align:center;
                padding:10px;
                border:1px solid gold;
            }

            td{
                background:#3E2723;
                color:white;
                text-align:center;
                padding:8px;
                border:1px solid #8D6E63;
            }

            </style>
            """,unsafe_allow_html=True)

            table="""

            <table>

            <tr>

            <th>م</th>

            <th>الرول</th>

            <th>تاريخ الجلسة</th>

            <th>الإجراءات</th>

            </tr>

            """

            for i,s in enumerate(sessions,start=1):

                table+=f"""

                <tr>

                <td>{i}</td>

                <td>{s[1]}</td>

                <td>{s[2]}</td>

                <td>{s[3]}</td>

                </tr>

                """

            table+="</table>"

            st.markdown(
                table,
                unsafe_allow_html=True
            )

        else:

            st.warning("لا توجد جلسات مسجلة")

        st.markdown("<br>",unsafe_allow_html=True)

        st.markdown("""
        <h3 style='color:#FFD700'>
        ⚖️ متابعة الجلسات
        </h3>
        """,unsafe_allow_html=True)

        c1,c2=st.columns(2)

        with c1:

            next_session_date=st.date_input(
                "تاريخ الجلسة القادمة"
            )

            next_roll=st.text_input(
                "الرول الجديد"
            )

        with c2:

            next_procedure=st.text_area(
                "سبب التأجيل / الإجراء",
                height=120
            )

        if st.button(
            "💾 حفظ الجلسة الجديدة",
            use_container_width=True
        ):

            cur.execute("""

            INSERT INTO sessions(

            case_id,
            session_date,
            roll_number,
            procedure,
            created_at

            )

            VALUES(

            ?,?,?,?,?

            )

            """,

            (

            case_id,

            str(next_session_date),

            next_roll,

            next_procedure,

            datetime.now().strftime(
                "%Y-%m-%d %H:%M:%S"
            )

            )

            )

            conn.commit()

            st.success("تم حفظ الجلسة بنجاح")

            st.rerun()
            # =====================================
        # تعديل جلسة سابقة
        # =====================================

        st.markdown("---")

        st.markdown("""
        <h3 style='color:#FFD700'>
        ✏️ تعديل جلسة
        </h3>
        """, unsafe_allow_html=True)

        cur.execute("""

        SELECT

        id,
        session_date

        FROM sessions

        WHERE case_id=?

        ORDER BY session_date ASC

        """, (case_id,))

        session_list = cur.fetchall()

        if session_list:

            session_names = [
                s[1] for s in session_list
            ]

            selected_date = st.selectbox(
                "اختر الجلسة",
                session_names
            )

            selected_session = None

            for s in session_list:

                if s[1] == selected_date:
                    selected_session = s[0]

            cur.execute("""

            SELECT

            roll_number,
            session_date,
            procedure

            FROM sessions

            WHERE id=?

            """, (selected_session,))

            data = cur.fetchone()

            edit_roll = st.text_input(
                "الرول",
                value=data[0]
            )

            edit_date = st.date_input(
                "تاريخ الجلسة",
                value=datetime.strptime(
                    data[1],
                    "%Y-%m-%d"
                )
            )

            edit_procedure = st.text_area(
                "الإجراءات",
                value=data[2],
                height=120
            )

            if st.button(
                "💾 حفظ التعديل",
                use_container_width=True
            ):

                cur.execute("""

                UPDATE sessions

                SET

                roll_number=?,
                session_date=?,
                procedure=?

                WHERE id=?

                """,

                (

                edit_roll,

                str(edit_date),

                edit_procedure,

                selected_session

                )

                )

                conn.commit()

                st.success("تم تعديل الجلسة بنجاح")

                st.rerun()

        else:

            st.info("لا توجد جلسات للتعديل")

        st.markdown("---")

        # =====================================
        # مستندات القضية
        # =====================================

        st.markdown("""
        <h3 style='color:#FFD700'>
        📂 مستندات القضية
        </h3>
        """, unsafe_allow_html=True)

        document_type = st.selectbox(

            "نوع المستند",

            [

                "صحيفة دعوى",

                "صحيفة استئناف",

                "صحيفة طعن",

                "مذكرة دفاع",

                "حافظة مستندات",

                "تقرير خبير",

                "تقرير طب شرعى",

                "تقرير لجنة طبية",

                "صحيفة تجديد من الشطب",

                "صحيفة تعجيل من الوقف",

                "صورة حكم تمهيدى",

                "أخرى"

            ]

        )

        document_description = st.text_input(
            "بيان المستند"
        )

        uploaded_document = st.file_uploader(
            "رفع المستند"
            )
        # =====================================
        # حفظ المستند
        # =====================================

        if st.button(
            "💾 حفظ المستند",
            use_container_width=True
        ):

            if uploaded_document is None:

                st.error("يرجى اختيار مستند")

            else:

                import os

                upload_folder = "documents"

                os.makedirs(
                    upload_folder,
                    exist_ok=True
                )

                file_path = os.path.join(
                    upload_folder,
                    uploaded_document.name
                )

                with open(
                    file_path,
                    "wb"
                ) as f:

                    f.write(
                        uploaded_document.getbuffer()
                    )

                cur.execute("""

                INSERT INTO documents(

                case_id,
                document_type,
                document_description,
                file_name,
                file_path,
                uploaded_at

                )

                VALUES(?,?,?,?,?,?)

                """,

                (

                case_id,

                document_type,

                document_description,

                uploaded_document.name,

                file_path,

                datetime.now().strftime(
                    "%Y-%m-%d %H:%M:%S"
                )

                )

                )

                conn.commit()

                st.success("تم حفظ المستند بنجاح")

                st.rerun()

        st.markdown("---")

        # =====================================
        # عرض مستندات القضية
        # =====================================

        st.markdown("""
        <h3 style='color:#FFD700'>
        📂 المستندات المحفوظة
        </h3>
        """, unsafe_allow_html=True)

        cur.execute("""

        SELECT

        id,
        document_type,
        document_description,
        file_name,
        file_path

        FROM documents

        WHERE case_id=?

        ORDER BY uploaded_at DESC

        """,(case_id,))

        docs = cur.fetchall()

        if docs:

            for doc in docs:

                with st.expander(
                    f"📄 {doc[1]}"
                ):

                    st.write(
                        f"**البيان :** {doc[2]}"
                    )

                    st.write(
                        f"**اسم الملف :** {doc[3]}"
                    )

                    try:

                        with open(
                            doc[4],
                            "rb"
                        ) as f:

                            st.download_button(

                                "⬇ تحميل المستند",

                                f,

                                file_name=doc[3],

                                use_container_width=True,

                                key=f"download_{doc[0]}"

                            )

                    except:

                        st.error(
                            "الملف غير موجود"
                        )

        else:

            st.info(
                "لا توجد مستندات"
            )

        st.markdown("---")

        # =====================================
        # جلسة الحكم
        # =====================================

        st.markdown("""
        <h3 style='color:#FFD700'>
        ⚖️ بيانات الحكم
        </h3>
        """, unsafe_allow_html=True)

        judgment_date = st.date_input(
            "تاريخ جلسة الحكم"
        )

        judgment_text = st.text_area(

            "منطوق الحكم",

            height=150

        )

        if st.button(

            "💾 حفظ الحكم",

            use_container_width=True

        ):

            cur.execute("""

            INSERT INTO sessions(

            case_id,

            session_date,

            roll_number,

            procedure,

            created_at

            )

            VALUES(

            ?,?,?,?,?

            )

            """,

            (

            case_id,

            str(judgment_date),

            "حكم",

            judgment_text,

            datetime.now().strftime(
                "%Y-%m-%d %H:%M:%S"
            )

            )

            )

            conn.commit()

            st.success(
                "تم حفظ بيانات الحكم"
            )

            st.rerun()

        st.markdown("<br>", unsafe_allow_html=True)

        b1, b2 = st.columns(2)

        with b1:

            if st.button(
                "⬅ العودة للحصر العام",
                use_container_width=True
            ):

                st.session_state.page = "inventory"

                st.rerun()
                # ==================== بداية الجزء الأول ====================

elif st.session_state.page == "reports":

    st.markdown("""
    <div style="
    background:#4E342E;
    border:3px solid #D4AF37;
    border-radius:18px;
    padding:18px;
    text-align:center;
    color:#FFD700;
    font-size:34px;
    font-weight:bold;">
    📊 التقارير
    </div>
    """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    report_type = st.selectbox(
        "نوع التقرير",
        [
            "بيان بجميع الدعاوى المتداولة",
            "بيان الدعاوى حسب موضوع الدعوى",

            "بيان بجميع الأحكام الصادرة",
            "بيان بالأحكام الصادرة للصالح",
            "بيان بالأحكام الصادرة للضد",
            "بيان الأحكام الصادرة حسب موضوع الدعوى",

            "التقارير العددية",

            "بيان عددى بجميع الدعاوى المتداولة",
            "بيان عددى بالدعاوى حسب موضوع الدعوى",

            "بيان عددى بجميع الأحكام الصادرة",
            "بيان عددى بالأحكام الصادرة للصالح",
            "بيان عددى بالأحكام الصادرة للضد",
            "بيان عددى بالأحكام الصادرة حسب موضوع الدعوى"
        ]
    )

    c1, c2 = st.columns(2)

    with c1:
        from_date = st.date_input("من تاريخ")

    with c2:
        to_date = st.date_input("إلى تاريخ")

    c3, c4 = st.columns(2)

    with c3:
        office = st.text_input(
            "ديوان عام منطقة",
            value=""
        )

    with c4:
        lawyer = st.text_input(
            "طرف الأستاذ / المحامى",
            value=""
        )

    subject_search = ""

    if report_type in (
        "بيان الدعاوى حسب موضوع الدعوى",
        "بيان الأحكام الصادرة حسب موضوع الدعوى",
        "بيان عددى بالدعاوى حسب موضوع الدعوى",
        "بيان عددى بالأحكام الصادرة حسب موضوع الدعوى"
    ):

        subject_search = st.text_input(
            "موضوع الدعوى"
        )

    st.markdown("---")

    show_report = st.button(
        "📄 عرض التقرير",
        use_container_width=True,
        type="primary"
    )

    if show_report:

        report_text = ""

        st.info("📄 التقرير")

        st.write("# الهيئة القومية للتأمين الاجتماعى")
        st.write("## الإدارة المركزية للشئون القانونية")
        st.write("## الإدارة العامة للقضايا")

        st.divider()

        a1, a2 = st.columns(2)

        with a1:
            st.write("**ديوان عام منطقة :**", office)

        with a2:
            st.write("**طرف الأستاذ / المحامى :**", lawyer)

        st.divider()

        st.write(f"## {report_type}")

        st.write(
            f"خلال الفترة من {from_date.strftime('%d/%m/%Y')} حتى {to_date.strftime('%d/%m/%Y')}"
        )

        st.divider()

# ==================== نهاية الجزء الأول ====================
# ==================================================
        # ==================================================
# بيان الدعاوى المتداولة
# ==================================================

        if report_type in (
            "بيان بجميع الدعاوى المتداولة",
            "بيان الدعاوى حسب موضوع الدعوى"
        ):

            sql = """
            SELECT
                c.case_number,
                c.judicial_year,
                c.circuit,
                c.case_category,
                c.court_type || ' ' || c.court_name,
                c.mission,
                c.plaintiff,
                c.defendant,
                c.subject,
                c.notes,
                (
                    SELECT session_date
                    FROM sessions s
                    WHERE s.case_id = c.id
                    ORDER BY session_date DESC,id DESC
                    LIMIT 1
                ) AS last_session,
                (
                    SELECT procedure
                    FROM sessions s
                    WHERE s.case_id = c.id
                    ORDER BY session_date DESC,id DESC
                    LIMIT 1
                ) AS last_procedure
            FROM cases c
            WHERE NOT EXISTS(
                SELECT 1
                FROM sessions s
                WHERE s.case_id=c.id
                AND s.roll_number='حكم'
            )
            """

            params = []

            if subject_search:

                sql += " AND c.subject LIKE ? "
                params.append(f"%{subject_search}%")

            sql += """
            ORDER BY
            last_session DESC,
            c.id DESC
            """

            cur.execute(sql, params)

            rows = cur.fetchall()

            if rows:

                data = []
                rows_word = []

                headers = [
                    "م",
                    "رقم القضية",
                    "السنة",
                    "الدائرة",
                    "نوع القضية",
                    "المحكمة",
                    "المأمورية",
                    "الخصوم",
                    "موضوع الدعوى",
                    "آخر جلسة",
                    "الإجراء",
                    "ملاحظات"
                ]

                serial = 1

                for r in rows:

                    row = {

                        "م": serial,

                        "رقم القضية": r[0],

                        "السنة": r[1],

                        "الدائرة": r[2],

                        "نوع القضية": r[3] if r[3] else "-",

                        "المحكمة": r[4],

                        "المأمورية": r[5] if r[5] else "-",

                        "الخصوم":
                        f"{r[6] if r[6] else '-'}\nضــــــد\n{r[7] if r[7] else '-'}",

                        "موضوع الدعوى": r[8] if r[8] else "-",

                        "آخر جلسة": r[10] if r[10] else "-",

                        "الإجراء": r[11] if r[11] else "-",

                        "ملاحظات": r[9] if r[9] else "-"

                    }

                    data.append(row)

                    rows_word.append([

                        row["م"],

                        row["رقم القضية"],

                        row["السنة"],

                        row["الدائرة"],

                        row["نوع القضية"],

                        row["المحكمة"],

                        row["المأمورية"],

                        row["الخصوم"],

                        row["موضوع الدعوى"],

                        row["آخر جلسة"],

                        row["الإجراء"],

                        row["ملاحظات"]

                    ])

                    serial += 1

                st.dataframe(
                    data,
                    use_container_width=True,
                    hide_index=True
                )
# ==================== بداية الجزء الثالث ====================

            else:

                st.warning("لا توجد دعاوى مطابقة لشروط البحث.")

        # ==================================================
        # بيان الأحكام الصادرة
        # ==================================================

        elif report_type in (
            "بيان بجميع الأحكام الصادرة",
            "بيان بالأحكام الصادرة للصالح",
            "بيان بالأحكام الصادرة للضد",
            "بيان الأحكام الصادرة حسب موضوع الدعوى"
        ):

            sql = """
            SELECT
                c.case_number,
                c.judicial_year,
                c.circuit,
                c.case_category,
                c.court_type || ' ' || c.court_name,
                c.mission,
                c.plaintiff,
                c.defendant,
                c.subject,
                c.notes,
                s.session_date,
                s.procedure
            FROM cases c
            INNER JOIN sessions s
            ON s.id = (
                SELECT id
                FROM sessions
                WHERE case_id = c.id
                AND roll_number='حكم'
                ORDER BY session_date DESC,id DESC
                LIMIT 1
            )
            WHERE 1=1
            """

            params = []

            if report_type == "بيان الأحكام الصادرة حسب موضوع الدعوى":
                if subject_search:
                    sql += " AND c.subject LIKE ? "
                    params.append(f"%{subject_search}%")

            sql += " ORDER BY s.session_date DESC "

            cur.execute(sql, params)

            rows = cur.fetchall()

            if rows:

                data = []
                rows_word = []
                serial = 1

                for r in rows:

                    procedure = str(r[11] or "").strip()

                    result = "غير محدد"

                    if "للصالح" in procedure:
                        result = "للصالح"

                    elif "للضد" in procedure:
                        result = "للضد"

                    if report_type == "بيان بالأحكام الصادرة للصالح":
                        if result != "للصالح":
                            continue

                    if report_type == "بيان بالأحكام الصادرة للضد":
                        if result != "للضد":
                            continue

                    row = {

                        "م": serial,
                        "رقم القضية": r[0],
                        "السنة": r[1],
                        "الدائرة": r[2],
                        "النوع": r[3],
                        "المحكمة": r[4],
                        "المأمورية": r[5] or "-",
                        "الخصوم": f"{r[6]}\nضد\n{r[7]}",
                        "موضوع الدعوى": r[8],
                        "تاريخ الحكم": r[10] or "-",
                        "منطوق الحكم": r[11] or "-",
                        "النتيجة": result,
                        "ملاحظات": r[9] or "-"
                    }

                    data.append(row)

                    rows_word.append([
                        row["م"],
                        row["رقم القضية"],
                        row["السنة"],
                        row["الدائرة"],
                        row["النوع"],
                        row["المحكمة"],
                        row["المأمورية"],
                        row["الخصوم"],
                        row["موضوع الدعوى"],
                        row["تاريخ الحكم"],
                        row["منطوق الحكم"],
                        row["النتيجة"],
                        row["ملاحظات"]
                    ])

                    serial += 1

# ==================== نهاية الجزء الثالث ====================
# ==================== بداية الجزء الرابع ====================

        elif report_type.startswith("بيان عددى"):

            sql = """
            SELECT
                c.subject,
                (
                    SELECT procedure
                    FROM sessions s
                    WHERE s.case_id = c.id
                    ORDER BY session_date DESC, id DESC
                    LIMIT 1
                ) AS procedure
            FROM cases c
            WHERE 1=1
            """

            params = []

            if report_type in (
                "بيان عددى بالدعاوى حسب موضوع الدعوى",
                "بيان عددى بالأحكام الصادرة حسب موضوع الدعوى"
            ):
                if subject_search:
                    sql += " AND c.subject LIKE ? "
                    params.append(f"%{subject_search}%")

            cur.execute(sql, params)

            rows = cur.fetchall()

            total = 0

            for subject, procedure in rows:

                procedure = str(procedure or "")

                if report_type == "بيان عددى بجميع الدعاوى المتداولة":

                    if "للصالح" not in procedure and "للضد" not in procedure:
                        total += 1

                elif report_type == "بيان عددى بالدعاوى حسب موضوع الدعوى":

                    if "للصالح" not in procedure and "للضد" not in procedure:
                        total += 1

                elif report_type == "بيان عددى بجميع الأحكام الصادرة":

                    if "للصالح" in procedure or "للضد" in procedure:
                        total += 1

                elif report_type == "بيان عددى بالأحكام الصادرة للصالح":

                    if "للصالح" in procedure:
                        total += 1

                elif report_type == "بيان عددى بالأحكام الصادرة للضد":

                    if "للضد" in procedure:
                        total += 1

                elif report_type == "بيان عددى بالأحكام الصادرة حسب موضوع الدعوى":

                    if "للصالح" in procedure or "للضد" in procedure:
                        total += 1

            data = [
                {
                    "البيان": report_type,
                    "العدد": total
                }
            ]

            headers = [
                "البيان",
                "العدد"
            ]

            rows_word = [
                [
                    report_type,
                    total
                ]
            ]

            st.dataframe(
                data,
                use_container_width=True,
                hide_index=True
            )

# ==================== نهاية الجزء الرابع ====================
# ==================== بداية الجزء الخامس ====================

        rows_word = locals().get("rows_word", [])
        headers = locals().get("headers", [])

        c1, c2, c3, c4 = st.columns(4)

        with c1:
            st.button(
                "🖨 طباعة",
                use_container_width=True,
                disabled=True
            )

        with c2:
            st.button(
                "📊 Excel",
                use_container_width=True,
                disabled=True
            )

        with c3:

            if rows_word:

                word_file = create_word_report(
                    report_title=report_type,
                    office=office,
                    lawyer=lawyer,
                    from_date=from_date,
                    to_date=to_date,
                    headers=headers,
                    rows=rows_word
                )

                st.download_button(
                    label="⬇ تحميل Word",
                    data=word_file,
                    file_name=f"{report_type}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

            else:

                st.button(
                    "⬇ تحميل Word",
                    use_container_width=True,
                    disabled=True
                )

        with c4:

            st.button(
                "⬇ تحميل PDF",
                use_container_width=True,
                disabled=True
            )

        st.divider()

        st.write("### وتفضلوا بقبول وافر الاحترام")

        s1, s2 = st.columns(2)

        with s1:
            st.write("عضو الإدارة")

        with s2:
            st.write("مدير الإدارة")

        st.write(
            f"تحريراً فى : {datetime.now().strftime('%d/%m/%Y')}"
        )

    st.markdown("<br>", unsafe_allow_html=True)

    if st.button(
        "⬅ العودة للرئيسية",
        use_container_width=True
    ):
        st.session_state.page = "home"
        st.rerun()

# ==================== نهاية الجزء الخامس ====================
